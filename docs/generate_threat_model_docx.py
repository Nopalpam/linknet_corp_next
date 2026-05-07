from pathlib import Path
import math
import zipfile

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "threat_modeling_assets"
ASSETS.mkdir(exist_ok=True)
OUT = DOCS / "Threat_Modeling_LinkNet_Corp_Next.docx"
FALLBACK_OUT = DOCS / "Threat_Modeling_LinkNet_Corp_Next_Revised_Post_Hardening.docx"

PROJECT = "LinkNet Corp Next"
TITLE = "Threat Modeling and Data Flow Diagram Security Assessment"
VERSION = "2.0 - Revised Post-Hardening"
DATE = "7 May 2026"
PREPARED_BY = "Codex - AI-assisted Security Review"
ORGANIZATION = "LinkNet Corp / PT Link Net Tbk"
CLASSIFICATION = "Internal Security Assessment"

NAVY = "17324D"
BLUE = "2563EB"
TEAL = "0F766E"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
PURPLE = "6D28D9"
GRAY = "374151"
LIGHT = "F8FAFC"
BORDER = "CBD5E1"
SOFT_BLUE = "EFF6FF"
SOFT_TEAL = "F0FDFA"
SOFT_GREEN = "F0FDF4"
SOFT_AMBER = "FFFBEB"
SOFT_RED = "FEF2F2"


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.replace("#", "")
    return RGBColor(int(value[:2], 16), int(value[2:4], 16), int(value[4:], 16))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, color: str = GRAY, size: int = 8) -> None:
    cell.text = ""
    lines = str(text).splitlines() or [""]
    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.font.color.rgb = rgb(color)
        run.bold = bold


def styled_table(doc, headers, rows, widths=None, style_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        shade(header_cells[idx], style_fill)
        border(header_cells[idx], style_fill, "8")
        set_cell_text(header_cells[idx], header, bold=True, color="FFFFFF", size=8)
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            header_cells[idx].width = Inches(widths[idx])
    repeat_header(table.rows[0])

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            shade(cells[col_idx], "FFFFFF" if row_idx % 2 == 0 else LIGHT)
            border(cells[col_idx])
            set_cell_text(cells[col_idx], value, size=8)
            cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[col_idx].width = Inches(widths[col_idx])
    doc.add_paragraph()
    return table


def add_info_box(doc, title: str, body: str, fill: str = SOFT_BLUE, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell, accent, "12")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(accent)
    for line in body.splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Aptos"
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(GRAY)
    doc.add_paragraph()


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_caption(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = rgb(GRAY)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_update_fields(doc) -> None:
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings._element.append(update)


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    add_field(paragraph, "PAGE", "1")
    paragraph.add_run(" of ")
    add_field(paragraph, "NUMPAGES", "1")


def configure_document(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = rgb(GRAY)
    normal.paragraph_format.space_after = Pt(6)

    for idx, (size, color) in enumerate(((17, NAVY), (13, BLUE), (10.5, TEAL)), start=1):
        style = styles[f"Heading {idx}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(12 if idx == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        if idx == 1:
            style.paragraph_format.keep_with_next = True

    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(28)
    styles["Title"].font.color.rgb = rgb(NAVY)
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8)
    styles["Caption"].font.color.rgb = rgb(GRAY)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{PROJECT} | {CLASSIFICATION}")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{TITLE} | ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(GRAY)

    set_update_fields(doc)


def load_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw, text: str, font, max_width: int):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def draw_box(draw, xy, title, subtitle="", fill="FFFFFF", outline=BORDER, text_color=NAVY, dashed=False):
    x1, y1, x2, y2 = xy
    if dashed:
        dash = 18
        for x in range(x1, x2, dash * 2):
            draw.line((x, y1, min(x + dash, x2), y1), fill=f"#{outline}", width=3)
            draw.line((x, y2, min(x + dash, x2), y2), fill=f"#{outline}", width=3)
        for y in range(y1, y2, dash * 2):
            draw.line((x1, y, x1, min(y + dash, y2)), fill=f"#{outline}", width=3)
            draw.line((x2, y, x2, min(y + dash, y2)), fill=f"#{outline}", width=3)
    else:
        draw.rounded_rectangle(xy, radius=22, fill=f"#{fill}", outline=f"#{outline}", width=3)
    title_font = load_font(25, bold=True)
    body_font = load_font(20)
    y = y1 + 24
    for line in wrap_text(draw, title, title_font, x2 - x1 - 40):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - bbox[2]) / 2, y), line, fill=f"#{text_color}", font=title_font)
        y += 30
    if subtitle:
        y += 8
        for line in wrap_text(draw, subtitle, body_font, x2 - x1 - 42):
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text(((x1 + x2 - bbox[2]) / 2, y), line, fill=f"#{GRAY}", font=body_font)
            y += 25


def draw_arrow(draw, start, end, label="", color=BLUE):
    draw.line((*start, *end), fill=f"#{color}", width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    left = (
        end[0] - length * math.cos(angle - math.pi / 6),
        end[1] - length * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - length * math.cos(angle + math.pi / 6),
        end[1] - length * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=f"#{color}")
    if label:
        font = load_font(18)
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        lines = wrap_text(draw, label, font, 210)
        label_h = 24 * len(lines) + 12
        label_w = max(draw.textbbox((0, 0), line, font=font)[2] for line in lines) + 22
        draw.rounded_rectangle((mx - label_w / 2, my - label_h / 2, mx + label_w / 2, my + label_h / 2), radius=10, fill="#FFFFFF", outline=f"#{BORDER}", width=1)
        y = my - label_h / 2 + 6
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((mx - bbox[2] / 2, y), line, fill=f"#{GRAY}", font=font)
            y += 24


def diagram_canvas(title: str):
    img = Image.new("RGB", (1800, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 96), fill=f"#{NAVY}")
    draw.text((46, 28), title, font=load_font(34, bold=True), fill="white")
    draw.text((46, 990), f"{PROJECT} | Revised post-hardening DFD | {DATE}", font=load_font(18), fill=f"#{GRAY}")
    return img, draw


def create_context_diagram():
    img, draw = diagram_canvas("Context Diagram")
    draw_box(draw, (720, 360, 1080, 570), "LinkNet Corp Next", "Admin CMS, public website, backend API, content and form services", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, (80, 180, 370, 330), "Public Visitors", "Website browsing, forms, downloads", fill=LIGHT, outline=GRAY)
    draw_box(draw, (80, 630, 370, 780), "CMS Admin Users", "Authenticated content and user management", fill=SOFT_AMBER, outline=AMBER)
    draw_box(draw, (1370, 150, 1670, 300), "PostgreSQL", "Prisma-managed application data", fill=SOFT_GREEN, outline=GREEN)
    draw_box(draw, (1370, 390, 1670, 555), "Object Storage", "S3 or Azure Blob; private by default", fill=SOFT_TEAL, outline=TEAL)
    draw_box(draw, (1370, 650, 1670, 835), "External Services", "GA4, LinkNet partner API, Yahoo finance, approved dispatch endpoints", fill="F5F3FF", outline=PURPLE)
    draw_arrow(draw, (370, 255), (720, 420), "HTTPS UI/API", BLUE)
    draw_arrow(draw, (370, 705), (720, 510), "HTTPS + HttpOnly cookies", BLUE)
    draw_arrow(draw, (1080, 420), (1370, 225), "SQL via Prisma", GREEN)
    draw_arrow(draw, (1080, 470), (1370, 470), "Validated files", TEAL)
    draw_arrow(draw, (1080, 525), (1370, 735), "Outbound HTTPS allowlisted", PURPLE)
    return save_diagram(img, "dfd_context.png")


def create_level0_diagram():
    img, draw = diagram_canvas("Level 0 DFD")
    draw_box(draw, (70, 205, 345, 360), "External Users", "Public visitors and CMS admins", fill=LIGHT, outline=GRAY)
    draw_box(draw, (465, 125, 770, 295), "Public Web", "Next.js web app and stock API routes", fill=SOFT_TEAL, outline=TEAL)
    draw_box(draw, (465, 405, 770, 585), "Admin Frontend", "Next.js admin UI; cookie session", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, (890, 240, 1185, 470), "Backend API", "Express middleware, validation, CSRF, rate limits, auth, RBAC", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, (1290, 115, 1625, 265), "CMS Services", "Pages, menus, news, settings, analytics", fill="FFFFFF", outline=BLUE)
    draw_box(draw, (1290, 335, 1625, 500), "Upload and Form Services", "MIME/ext/magic scanner, dispatch allowlist", fill="FFFFFF", outline=TEAL)
    draw_box(draw, (1290, 590, 1625, 760), "Data Stores", "PostgreSQL, S3/Azure Blob/local uploads", fill=SOFT_GREEN, outline=GREEN)
    draw_box(draw, (1290, 820, 1625, 940), "External APIs", "GA4, stock provider, partner API, approved webhooks", fill="F5F3FF", outline=PURPLE)
    draw_arrow(draw, (345, 270), (465, 215), "Public HTTPS", TEAL)
    draw_arrow(draw, (345, 295), (465, 495), "Admin HTTPS", BLUE)
    draw_arrow(draw, (770, 495), (890, 355), "API requests + CSRF", BLUE)
    draw_arrow(draw, (770, 215), (890, 330), "Public API calls", TEAL)
    draw_arrow(draw, (1185, 300), (1290, 190), "Authorized service calls", BLUE)
    draw_arrow(draw, (1185, 385), (1290, 420), "Validated upload/form flow", TEAL)
    draw_arrow(draw, (1460, 500), (1460, 590), "Persistence", GREEN)
    draw_arrow(draw, (1185, 430), (1290, 880), "Outbound HTTPS", PURPLE)
    return save_diagram(img, "dfd_level0.png")


def create_auth_diagram():
    img, draw = diagram_canvas("Level 1 DFD - Authentication and Authorization")
    boxes = {
        "browser": (65, 280, 350, 455),
        "frontend": (460, 280, 760, 455),
        "auth": (875, 155, 1195, 335),
        "middleware": (875, 445, 1195, 640),
        "db": (1315, 150, 1635, 325),
        "protected": (1315, 445, 1635, 640),
        "audit": (1315, 730, 1635, 895),
    }
    draw_box(draw, boxes["browser"], "Admin Browser", "Login, MFA, protected CMS actions", fill=LIGHT, outline=GRAY)
    draw_box(draw, boxes["frontend"], "Admin Frontend", "No bearer token in localStorage; credentials include", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, boxes["auth"], "Auth Controller", "bcrypt password check, refresh rotation, MFA verify when required", fill=SOFT_AMBER, outline=AMBER)
    draw_box(draw, boxes["middleware"], "Security Middleware", "HttpOnly cookie auth, CSRF, rate limit, RBAC permissions", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, boxes["db"], "Auth Data Store", "users, roles, permissions, refresh token hashes, MFA secret", fill=SOFT_GREEN, outline=GREEN)
    draw_box(draw, boxes["protected"], "Protected APIs", "CMS, user, role, upload, dashboard, analytics APIs", fill="FFFFFF", outline=BLUE)
    draw_box(draw, boxes["audit"], "Audit Logging", "Actor, route, action, timestamp and sanitized metadata", fill="FFFFFF", outline=GRAY)
    draw_arrow(draw, (350, 360), (460, 360), "HTTPS login", BLUE)
    draw_arrow(draw, (760, 330), (875, 245), "POST /auth/login", BLUE)
    draw_arrow(draw, (1195, 230), (1315, 225), "verify/hash lookup", GREEN)
    draw_arrow(draw, (1195, 305), (1315, 260), "persist refresh hash", GREEN)
    draw_arrow(draw, (875, 535), (760, 420), "Set HttpOnly cookies + CSRF", AMBER)
    draw_arrow(draw, (760, 410), (875, 535), "cookie session + X-CSRF-Token", BLUE)
    draw_arrow(draw, (1195, 535), (1315, 535), "permission-gated request", BLUE)
    draw_arrow(draw, (1480, 640), (1480, 730), "activity event", GRAY)
    return save_diagram(img, "dfd_auth_level1.png")


def create_upload_diagram():
    img, draw = diagram_canvas("Level 1 DFD - Upload, Storage and Dispatch")
    draw_box(draw, (65, 210, 345, 380), "Uploader", "Admin CMS or public form submitter", fill=LIGHT, outline=GRAY)
    draw_box(draw, (455, 150, 760, 340), "Upload Routes", "uploadRateLimiter, public form limiter, route auth/RBAC when applicable", fill=SOFT_BLUE, outline=BLUE)
    draw_box(draw, (455, 490, 760, 675), "Presigned Upload", "Disabled in production by default; opt-in only", fill=SOFT_AMBER, outline=AMBER)
    draw_box(draw, (890, 145, 1195, 340), "Validation Pipeline", "size, extension, MIME, magic bytes, filename safety", fill=SOFT_TEAL, outline=TEAL)
    draw_box(draw, (890, 485, 1195, 675), "Quarantine Path", "direct object upload must be verified before public use", fill=SOFT_AMBER, outline=AMBER)
    draw_box(draw, (1320, 135, 1635, 315), "Application Storage", "local disk, S3, or Azure Blob; private by default", fill=SOFT_GREEN, outline=GREEN)
    draw_box(draw, (1320, 445, 1635, 620), "PostgreSQL Metadata", "file metadata, form submissions, ownership and audit trail", fill=SOFT_GREEN, outline=GREEN)
    draw_box(draw, (1320, 760, 1635, 915), "Approved Dispatch", "HTTPS endpoint allowlist; private hosts blocked", fill="F5F3FF", outline=PURPLE)
    draw_arrow(draw, (345, 295), (455, 245), "multipart upload", BLUE)
    draw_arrow(draw, (760, 245), (890, 245), "validated request", TEAL)
    draw_arrow(draw, (1195, 245), (1320, 225), "store safe object", GREEN)
    draw_arrow(draw, (1040, 340), (1480, 445), "write metadata", GREEN)
    draw_arrow(draw, (345, 325), (455, 585), "optional direct flow", AMBER)
    draw_arrow(draw, (760, 585), (890, 585), "short TTL + quarantine", AMBER)
    draw_arrow(draw, (1195, 585), (1320, 230), "private object write", GREEN)
    draw_arrow(draw, (1040, 675), (1320, 840), "approved outbound only", PURPLE)
    return save_diagram(img, "dfd_upload_level1.png")


def save_diagram(img, filename: str) -> Path:
    path = ASSETS / filename
    img.save(path, quality=96)
    return path


def add_cover(doc) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(CLASSIFICATION)
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(BLUE)
    r.bold = True

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(PROJECT)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(TITLE)
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.color.rgb = rgb(NAVY)
    r.bold = True

    p = doc.add_paragraph()
    r = p.add_run("Revised enterprise threat model reflecting final application hardening, residual risk validation, and actual implementation evidence.")
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(GRAY)

    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Version", VERSION),
        ("Prepared By", PREPARED_BY),
        ("Date", DATE),
        ("Company / Organization", ORGANIZATION),
        ("Assessment Scope", "Source code, configuration, middleware, authentication, API routes, upload/storage, Docker/Kubernetes references, and existing threat model document"),
        ("Final Risk Posture", "Critical: 0 | High: 0 | Medium: minimized with mitigation | Low: acceptable"),
    ]
    for row_idx, (key, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        shade(cells[0], NAVY)
        shade(cells[1], "FFFFFF" if row_idx % 2 == 0 else LIGHT)
        border(cells[0], NAVY, "8")
        border(cells[1])
        set_cell_text(cells[0], key, True, "FFFFFF", 9)
        set_cell_text(cells[1], value, False, GRAY, 9)
    doc.add_paragraph()
    add_info_box(
        doc,
        "Assessment Position",
        "This document intentionally avoids theoretical findings that are not evidenced in the current codebase. Severity reflects current compensating controls and exploitability after remediation.",
        SOFT_GREEN,
        GREEN,
    )
    doc.add_page_break()


def add_toc(doc) -> None:
    add_heading(doc, "Table of Contents", 1)
    paragraph = doc.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Table of contents will update automatically in Microsoft Word.")
    doc.add_page_break()


def paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(GRAY)


def add_bullets(doc, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(GRAY)


def build_document():
    diagrams = [
        ("Figure 1. Context Diagram", create_context_diagram()),
        ("Figure 2. Level 0 Data Flow Diagram", create_level0_diagram()),
        ("Figure 3. Level 1 Authentication and Authorization Flow", create_auth_diagram()),
        ("Figure 4. Level 1 Upload, Storage and Dispatch Flow", create_upload_diagram()),
    ]

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_info_box(
        doc,
        "Final Security Posture",
        "Critical Risk: 0\nHigh Risk: 0\nMedium Risk: limited to residual items with active mitigations and documented next-step hardening\nLow Risk: acceptable and monitored",
        SOFT_GREEN,
        GREEN,
    )
    paragraph(
        doc,
        "The LinkNet Corp Next codebase was reviewed and remediated across authentication, authorization, upload handling, storage access, operational endpoints, external integration controls, dependency posture, and environment/secret handling. The revised assessment reflects the hardened implementation rather than the earlier document's pre-remediation risk state.",
    )
    paragraph(
        doc,
        "No residual Critical or High risks were identified after remediation. Remaining Medium risks are realistic residual concerns tied to operational hardening, dependency lifecycle, encryption depth, or defense-in-depth controls. Each Medium item has existing mitigation and recommended follow-up.",
    )
    styled_table(
        doc,
        ["Risk Tier", "Residual Count", "Assessment Result"],
        [
            ["Critical", "0", "No active Critical risk remains in the reviewed codebase."],
            ["High", "0", "Previously high-risk items were remediated or reclassified based on implemented controls."],
            ["Medium", "5", "Residual risk remains controlled by authentication, RBAC, validation, rate limiting, private storage defaults, allowlists, or audit controls."],
            ["Low", "7", "Accepted operational or improvement items with low exploitability."],
        ],
        [1.4, 1.3, 4.8],
        GREEN,
    )

    add_heading(doc, "2. Document Information", 1)
    styled_table(
        doc,
        ["Field", "Value"],
        [
            ["Project Name", PROJECT],
            ["Document Title", TITLE],
            ["Version", VERSION],
            ["Date", DATE],
            ["Prepared By", PREPARED_BY],
            ["Organization", ORGANIZATION],
            ["Source of Truth", "Current source code and configuration under C:/wamp64/www/linknet_corp_next plus the previous DOCX threat modeling artifact."],
            ["Assessment Standard", "STRIDE-oriented threat modeling with realistic post-control residual risk classification."],
        ],
        [2.2, 5.2],
    )

    add_heading(doc, "3. System Overview", 1)
    paragraph(
        doc,
        "LinkNet Corp Next is a monorepo containing a public Next.js website, a Next.js administrative frontend, an Express.js/TypeScript backend API, Prisma-managed PostgreSQL persistence, upload/file handling services, and optional integrations with S3, Azure Blob Storage, Azure Key Vault, Google Analytics Data API, partner APIs, and stock quote APIs.",
    )
    add_heading(doc, "3.1 High Level Architecture", 2)
    add_bullets(
        doc,
        [
            "Public website serves content and selected public API flows, including stock quote/historical endpoints with validation and IP rate limiting.",
            "Admin frontend uses cookie-based sessions and CSRF token forwarding; bearer tokens are not persisted in browser localStorage in production mode.",
            "Backend API centralizes authentication, authorization, RBAC, security headers, rate limiting, CSRF checks, upload validation, audit logging, and business/domain services.",
            "PostgreSQL stores users, roles, permissions, content, form submissions, file metadata, refresh token hashes, and audit/activity data through Prisma.",
            "Object storage is supported through local storage, AWS S3, or Azure Blob. Public ACL/public blob access is disabled by default and requires explicit opt-in.",
        ],
    )
    add_heading(doc, "3.2 Main Business Flow", 2)
    styled_table(
        doc,
        ["Flow", "Implementation Evidence", "Security-Relevant Controls"],
        [
            ["CMS administration", "Admin Next.js frontend to Express API routes for users, roles, pages, menus, news, settings, dashboard and analytics.", "HttpOnly cookies, CSRF token, auth middleware, permission middleware, rate limits and audit logging."],
            ["Public content delivery", "Public Next.js web app and backend public/content routes.", "Public-only routes separated from admin routes; CSP and input validation applied."],
            ["Form submission", "Form module routes and dispatch service.", "Public submission rate limit, validated upload path, HTTPS dispatch endpoint policy, production allowlist and private host blocking."],
            ["File upload", "Upload routes, filemanager routes and form file routes.", "File size limits, extension/MIME validation, magic-byte scanning, upload rate limit, private storage default and quarantine for presigned flow."],
            ["Operational health", "Public /health and /ready; detailed diagnostics protected in production.", "Detailed diagnostics require health token or private internal diagnostic allowance."],
        ],
        [1.4, 3.1, 3.0],
    )
    add_heading(doc, "3.3 Main Actors and Users", 2)
    styled_table(
        doc,
        ["Actor", "Trust Level", "Primary Actions"],
        [
            ["Public Visitor", "Untrusted external entity", "Browse public website, submit public forms, request public content and stock information."],
            ["CMS Admin User", "Authenticated internal/business user", "Manage CMS content, files, users, roles, settings, dashboard and analytics based on assigned permissions."],
            ["System Operator", "Privileged operational actor", "Configure runtime secrets, storage credentials, Kubernetes/infra settings and health monitoring."],
            ["External Service", "Third-party or partner boundary", "Receive approved dispatch calls or provide analytics, stock, media or partner data."],
        ],
        [1.8, 1.8, 3.8],
    )

    add_heading(doc, "4. Application Components", 1)
    styled_table(
        doc,
        ["Component", "Technology / Implementation", "Security Notes"],
        [
            ["Frontend - Admin", "Next.js 16, React 19, TypeScript, admin service layer.", "Cookie-only session flow; no production bearer token localStorage usage; CSRF header sourced from cookie."],
            ["Frontend - Public Web", "Next.js web application with public API routes.", "Stock APIs now validate symbols/interval/date range and apply per-IP rate limiting."],
            ["Backend/API", "Node.js, Express.js, TypeScript.", "Helmet/security headers, CORS, auth middleware, CSRF, RBAC permissions, rate limiters, validators, upload scanner and centralized error handling."],
            ["Database", "PostgreSQL via Prisma.", "Stores users, RBAC, refresh token hashes, CMS content, form submissions, files and audit activity."],
            ["Storage", "Local uploads, AWS S3, Azure Blob Storage.", "S3 public ACL and Azure public blob access are opt-in only; default posture is private storage."],
            ["Authentication Provider", "Internal auth service using bcrypt, JWT utilities, HttpOnly cookies, refresh token hashing and optional MFA.", "Token return in JSON body disabled unless AUTH_RETURN_TOKENS_IN_BODY=true; production mock auth disabled."],
            ["Cloud/Infrastructure", "Dockerfile/Kubernetes manifests, Azure Key Vault CSI references.", "Runtime secrets expected through environment/secret providers; tracked .env files removed from Git index and ignored going forward."],
            ["Queue / Message Broker", "Bull/ioredis dependencies present.", "Active production queue topology Not Found in Current Codebase."],
            ["Payment Gateway", "Not Found in Current Codebase.", "No payment flow assessed."],
        ],
        [1.7, 2.5, 3.3],
    )

    add_heading(doc, "4.1 API and Route Inventory", 2)
    styled_table(
        doc,
        ["API / Route Group", "Purpose", "Primary Controls"],
        [
            ["/api/auth/*", "Login, refresh, logout, password and MFA flows.", "Rate limiting, bcrypt password verification, refresh token hashing, HttpOnly cookies, CSRF-aware session handling."],
            ["/api/users, /api/roles, /api/profile", "Identity, profile and role management.", "Authentication, RBAC permissions, validators and audit logging."],
            ["/api/cms/pages, /api/news, /api/menu, /api/settings", "CMS content and configuration.", "Authentication plus explicit permission checks for read/create/update/delete where implemented."],
            ["/api/upload, /api/filemanager, /api/files", "Upload, file metadata and media handling.", "Upload rate limits, size/MIME/extension/magic-byte validation, private storage defaults and RBAC where applicable."],
            ["/api/form-modules/*", "Form rendering, submission and dispatch.", "Public submission rate limit, upload scanner and outbound endpoint policy."],
            ["/api/analytics, /api/dashboard", "Dashboard and analytics operations.", "Authentication and explicit operational/reporting permissions."],
            ["/health, /ready, /env-check, /health/detailed", "Health and diagnostics.", "Basic probes public; detailed diagnostics protected in production by token or private internal policy."],
            ["/api/stock/quote, /api/stock/historical", "Public stock data proxy in web app.", "Input validation and per-IP rate limits."],
        ],
        [2.0, 2.5, 3.1],
    )

    add_heading(doc, "4.2 Data Store Summary", 2)
    styled_table(
        doc,
        ["Data Store", "Representative Records", "Security Classification"],
        [
            ["PostgreSQL - users", "Email, username, hashed password, status, lockout, MFA state and MFA secret.", "Sensitive identity data; protected by DB boundary, RBAC and application access controls."],
            ["PostgreSQL - RBAC", "roles, permissions, role_permissions, user_roles.", "Authorization-critical configuration."],
            ["PostgreSQL - refresh_tokens", "Refresh token IDs and hashes with expiration.", "Session-sensitive; token material stored as hashes."],
            ["PostgreSQL - CMS/content", "Pages, menus, news, settings and metadata.", "Business content; write operations permission-gated."],
            ["PostgreSQL - forms/files/logs", "Form submissions, file metadata, activity logs.", "Potential PII and audit data; protected by RBAC and logging redaction."],
            ["Object storage", "Uploaded documents/images/media.", "Sensitive depending on upload context; private by default with explicit opt-in for public storage behavior."],
        ],
        [1.8, 3.1, 2.6],
    )

    add_heading(doc, "5. Security Remediation Performed", 1)
    styled_table(
        doc,
        ["Area", "Remediation Implemented", "Final Severity Impact"],
        [
            ["Secret exposure", ".env files removed from Git index and nested .env ignore rules added while preserving .env.example templates.", "Previous High/Critical-class concern reduced; no residual High when repository policy is followed."],
            ["Authentication/session", "Backend defaults to HttpOnly cookies only; JSON token body disabled by default; frontend no longer persists bearer tokens in localStorage for production auth.", "Session theft exposure reduced from High to Low/Medium residual browser-hardening concern."],
            ["Authorization", "CMS page routes, dashboard routes and analytics routes now have explicit permission middleware; existing controller-level RBAC reviewed for form modules.", "Unauthorized admin function access reduced to Low residual configuration risk."],
            ["Upload handling", "File size validation now covers req.file and req.files; scanner applied to upload, filemanager and form upload paths; upload-specific rate limits added.", "Upload abuse risk reduced from High to Medium residual AV/quarantine maturity risk."],
            ["Presigned uploads", "Production presigned uploads disabled by default; opt-in requires validated filename/content type/size, short TTL and quarantine prefix.", "Direct object upload risk reduced from High to Low/Medium depending on deployment enablement."],
            ["Storage access", "S3 public-read ACL and Azure public blob access disabled by default; explicit env opt-in required.", "Bucket/blob exposure risk reduced from High to Low residual configuration risk."],
            ["Operational endpoints", "Detailed health and env diagnostics protected in production by token or private internal diagnostics policy.", "Information disclosure risk reduced to Low."],
            ["External dispatch", "Form dispatch endpoint validation enforces HTTPS, blocks localhost/private IP targets and requires production allowlist.", "SSRF/data exfiltration risk reduced from High to Medium residual allowlist/egress governance risk."],
            ["Dependency posture", "npm audit fixes applied across frontend/web/filemanager/backend where safe; current high/critical advisories eliminated.", "Dependency risk reduced to Medium residual lifecycle tracking."],
            ["Standalone filemanager", "API key is fail-closed in production; wildcard CORS rejected in production.", "Unauthenticated file manager exposure reduced to Low."],
        ],
        [1.8, 4.3, 1.5],
        TEAL,
    )

    add_heading(doc, "6. Data Flow Diagrams", 1)
    add_info_box(
        doc,
        "Diagram Scope",
        "The diagrams below model the actual implementation boundaries found in the repository after remediation. They intentionally separate public flows, admin flows, protected backend middleware, database, storage and external-service trust boundaries.",
        SOFT_BLUE,
        BLUE,
    )
    for caption, path in diagrams:
        doc.add_picture(str(path), width=Inches(7.3))
        add_caption(doc, caption)

    add_heading(doc, "7. Trust Boundary Identification", 1)
    styled_table(
        doc,
        ["Boundary", "Components Inside Boundary", "Cross-Boundary Controls"],
        [
            ["Public zone", "Public visitor browser, public website and public stock/form endpoints.", "HTTPS, rate limits, strict input validation, public route segregation and safe error handling."],
            ["Admin zone", "Admin browser and Next.js admin UI.", "Authenticated cookie session, CSRF token, production auth-enabled guard and no localStorage bearer-token storage."],
            ["Backend internal zone", "Express middleware, controllers, services, upload scanner and dispatch services.", "Centralized security middleware, route authorization, validation, audit logging and redacted errors."],
            ["Database boundary", "PostgreSQL accessed through Prisma.", "Application-only DB access, hashed passwords/tokens, RBAC-controlled data access and migration/secret configuration."],
            ["Cloud/storage boundary", "S3, Azure Blob and local upload path.", "Private default storage, no public ACL/blob unless explicitly opted in, upload validation and presigned quarantine."],
            ["External service boundary", "GA4, stock provider, LinkNet partner API and form dispatch endpoints.", "HTTPS-only outbound flow, production allowlist, credential env management and sanitized logging."],
            ["Operational boundary", "Kubernetes/Docker/runtime secrets and health monitoring.", "Secret injection expected via env/Key Vault; detailed diagnostics protected in production."],
        ],
        [1.8, 2.6, 3.1],
    )

    add_heading(doc, "8. Existing Security Controls", 1)
    styled_table(
        doc,
        ["Control Family", "Implemented Control", "Evidence / Notes"],
        [
            ["Authentication", "JWT utilities, HttpOnly cookie session, refresh token rotation/hash storage and optional MFA.", "auth.controller.ts, mfa.controller.ts, authResponse.util.ts, jwt utilities, Prisma RefreshToken model."],
            ["Authorization", "Role/permission model and route-level permission checks for admin APIs.", "roles, permissions, user_roles and route middleware such as cms/page.routes.ts, dashboard.routes.ts and analytics.routes.ts."],
            ["CSRF", "CSRF middleware and frontend X-CSRF-Token forwarding with credentials include.", "Backend csrf.middleware.ts and frontend base service layer."],
            ["Rate limiting", "Login/auth, public form submission, upload and public stock API rate limiters.", "rateLimiter.middleware.ts and web/lib/apiRateLimit.ts."],
            ["Upload validation", "File size, extension, MIME and magic-byte validation.", "upload.middleware.ts, file-upload-scanner service and upload/filemanager/form routes."],
            ["Secure headers", "Helmet/security header middleware and Next.js CSP.", "securityHeaders.middleware.ts and frontend/web next.config files."],
            ["Secret handling", "Environment validator and .env Git ignore/de-indexing.", ".gitignore, environmentValidator.ts, backend/.env.example."],
            ["Storage security", "Private-by-default S3/Azure behavior and explicit opt-in for public access.", "s3Service.ts and azureStorage.service.ts."],
            ["Logging", "Activity logging and external API log redaction.", "activityLogger middleware and linknetEnterprise controller redaction."],
            ["Operational diagnostics", "Detailed health/env checks protected in production.", "health.routes.ts."],
        ],
        [1.7, 3.2, 2.6],
    )

    add_heading(doc, "9. STRIDE Threat Analysis", 1)
    styled_table(
        doc,
        ["STRIDE", "Threat Description", "Affected Component", "Existing Mitigation", "Residual Risk", "Recommended Mitigation"],
        [
            ["Spoofing", "Attacker attempts to impersonate an admin or reuse stolen tokens.", "Admin auth/session", "bcrypt password verification, lockout fields, optional MFA, HttpOnly cookies, refresh token hashes and no frontend localStorage bearer token storage.", "Medium", "Consider mandatory MFA for privileged roles and risk-based login alerts."],
            ["Spoofing", "Unauthenticated access to file manager in production.", "Filemanager service", "Production fails closed when API_KEY is missing; wildcard CORS rejected.", "Low", "Use centralized identity or mTLS/API gateway for production file manager access."],
            ["Tampering", "Unauthorized content update through admin CMS APIs.", "CMS pages/news/menu/settings", "Authentication and permission middleware; page routes now include explicit CRUD permission checks.", "Low", "Periodic RBAC route coverage tests and least-privilege role review."],
            ["Tampering", "Malicious or malformed upload attempts.", "Upload/filemanager/form upload routes", "Upload rate limit, size validation, MIME/extension checks, magic-byte scanner and storage quarantine for presigned flow.", "Medium", "Add production malware scanning engine and asynchronous quarantine release workflow."],
            ["Repudiation", "Admin denies sensitive changes.", "Admin APIs and CMS operations", "Activity logging captures actor, route, action and timestamp with sanitized metadata.", "Low", "Forward immutable audit logs to SIEM/WORM storage."],
            ["Information Disclosure", "Detailed environment or health data exposed publicly.", "Operational health endpoints", "Only /health and /ready remain public; /env-check and /health/detailed protected in production.", "Low", "Restrict diagnostic endpoints at ingress and monitoring network layer."],
            ["Information Disclosure", "Secrets or credentials committed to repository.", "Source control and env config", ".env files removed from Git index; nested ignore rules added; .env.example maintained.", "Low", "Run secret scanning in CI and rotate any secrets previously committed."],
            ["Information Disclosure", "External API credentials leaked through error logs.", "Partner/media/dispatch integrations", "External partner errors redacted; media token salt moved to env with production requirement.", "Low", "Centralized structured log redaction and DLP rules."],
            ["Denial of Service", "Abuse of public form, upload or stock API endpoints.", "Public endpoints and upload paths", "Public form limiter, upload limiter, symbol/date validation and per-IP stock API rate limits.", "Medium", "Add edge/WAF rate limits and request body size enforcement at reverse proxy."],
            ["Elevation of Privilege", "Authenticated user invokes admin function beyond role.", "Admin backend APIs", "Permission middleware and Prisma RBAC model; newly added checks on page/dashboard/analytics routes.", "Low", "Automated route-to-permission regression tests and admin role review cadence."],
            ["Tampering / SSRF", "Form dispatch endpoint used to target internal services.", "FormSubmissionDispatch service", "HTTPS required, localhost/private hosts blocked and FORM_DISPATCH_ALLOWED_HOSTS required in production.", "Medium", "Add network egress firewall and DNS pinning/egress proxy policy."],
            ["Information Disclosure", "Uploaded object accidentally exposed publicly.", "S3/Azure Blob storage", "S3 public ACL and Azure public blob access disabled by default; explicit opt-in required.", "Low", "Enforce bucket policies, block-public-access and cloud posture scanning."],
        ],
        [0.9, 2.0, 1.5, 2.3, 0.8, 2.0],
        NAVY,
    )

    add_heading(doc, "10. Authentication and Authorization Security", 1)
    add_heading(doc, "10.1 Login and Session Flow", 2)
    styled_table(
        doc,
        ["Step", "Implementation", "Security Assessment"],
        [
            ["Credential submission", "Admin frontend submits login to backend auth route over HTTPS in deployment context.", "Rate limits and validation reduce brute force and malformed request risk."],
            ["Credential verification", "Backend verifies password using bcrypt-backed password hashes and account status/lockout fields.", "Password exposure risk is controlled at rest through hashing; lockout telemetry exists."],
            ["MFA handling", "MFA verification controller exists and returns session using the same hardened response helper.", "MFA capability implemented; mandatory MFA for all privileged roles is Not Found in Current Codebase."],
            ["Session issue", "Backend sets HttpOnly cookies and CSRF cookie; token body return disabled unless AUTH_RETURN_TOKENS_IN_BODY=true.", "Browser JavaScript token theft risk significantly reduced."],
            ["Session use", "Frontend service layer sends credentials include and X-CSRF-Token; Authorization bearer header removed for production auth flow.", "CSRF middleware has meaningful protection because session cookies are the auth transport."],
            ["Refresh/logout", "Refresh uses cookie-based flow; refresh token hashes stored server-side; logout clears session cookie path.", "Refresh token database compromise requires hash cracking and valid token context."],
        ],
        [1.4, 3.3, 2.8],
    )

    add_heading(doc, "10.2 Access Control", 2)
    paragraph(
        doc,
        "Access control is implemented through authentication middleware and role/permission checks. The reviewed remediation added explicit permission checks to previously weaker CMS page, dashboard and analytics route groups. Form module CMS handlers were reviewed and found to already include controller-level permission arrays.",
    )
    add_info_box(
        doc,
        "Not Found in Current Codebase",
        "No external enterprise SSO/OAuth identity provider is enforced in the reviewed codebase. Authentication is implemented internally. Mandatory MFA for all privileged users is also not enforced globally in the reviewed code.",
        SOFT_AMBER,
        AMBER,
    )

    add_heading(doc, "11. Sensitive Data Handling", 1)
    styled_table(
        doc,
        ["Sensitive Data", "Current Handling", "Residual Risk", "Recommendation"],
        [
            ["Passwords", "Stored as bcrypt hashes; seed password requires SEED_ADMIN_PASSWORD in production.", "Low", "Continue enforcing password policy and rotate any legacy/default credentials."],
            ["Refresh tokens", "Stored as token hashes with expiry; session delivered through HttpOnly cookies.", "Low", "Monitor refresh token reuse and add anomaly detection."],
            ["MFA secrets", "Stored in user records.", "Medium", "Encrypt MFA secrets at field level using KMS/Key Vault envelope encryption."],
            ["Personal data and form submissions", "Persisted in PostgreSQL and controlled through app/RBAC boundaries.", "Medium", "Add retention policy, DLP classification and field-level encryption for high-sensitivity fields."],
            ["Upload files", "Validated by size, extension, MIME and magic bytes before storage; presigned uploads quarantine by default.", "Medium", "Add malware scanning engine and controlled release workflow."],
            ["Environment variables", ".env files ignored/de-indexed; examples retained; environment validator checks production-critical settings.", "Low", "Use CI secret scanning and managed secret stores only."],
            ["External API credentials", "Loaded from environment; logs redacted for partner API failure paths.", "Low", "Rotate secrets and enforce managed identity where available."],
        ],
        [1.7, 3.0, 1.0, 2.0],
        TEAL,
    )

    add_heading(doc, "12. External Integration and Infrastructure Risk", 1)
    styled_table(
        doc,
        ["Integration", "Implementation Found", "Risk Assessment", "Control / Mitigation"],
        [
            ["AWS S3", "AWS SDK and S3 service implementation.", "Low residual storage exposure risk.", "Public ACL disabled by default; explicit S3_ALLOW_PUBLIC_ACL=true required for public-read behavior."],
            ["Azure Blob Storage", "Azure storage service implementation.", "Low residual storage exposure risk.", "Container public access disabled by default; explicit AZURE_BLOB_PUBLIC_ACCESS=true required."],
            ["Azure Key Vault", "Azure identity/keyvault dependencies and Kubernetes CSI manifest references.", "Low configuration risk.", "Runtime secret injection expected; no plain secret material should be stored in source."],
            ["Google Analytics Data API", "Backend dependency and analytics route group.", "Low/Medium operational risk.", "Analytics routes require authentication and permissions."],
            ["Yahoo/stock provider", "Public web API proxy routes.", "Medium DoS/cost/external dependency risk.", "Symbol/date validation and per-IP rate limits implemented."],
            ["LinkNet partner/media APIs", "Controller-level external fetch flows.", "Low residual secret/log risk.", "External error logging redacted; token salt moved to env."],
            ["Form dispatch/webhook", "Outbound dispatch service for form submissions.", "Medium SSRF/misconfiguration residual risk.", "HTTPS-only, localhost/private host blocked and production allowlist required."],
            ["Email service", "Not Found in Current Codebase as a configured provider.", "Not assessed.", "Document once implemented."],
            ["Payment gateway", "Not Found in Current Codebase.", "Not assessed.", "Document once implemented."],
            ["Queue/message broker", "Bull/ioredis dependencies found; active production topology Not Found in Current Codebase.", "Not assessed as active flow.", "Threat model separately if queue processing is enabled."],
        ],
        [1.7, 2.3, 1.5, 2.2],
    )

    add_heading(doc, "13. Final Risk Assessment", 1)
    styled_table(
        doc,
        ["ID", "Residual Risk", "Severity", "Justification", "Existing Mitigation", "Recommended Next Step"],
        [
            ["M-01", "CSP still relies on limited unsafe-inline allowances for frontend/admin runtime compatibility.", "Medium", "Browser script injection impact can be meaningful, but token theft impact is reduced because auth tokens are not in localStorage and unsafe-eval/http development allowances are removed in production.", "HttpOnly cookies, CSRF, production CSP hardening and no localStorage bearer token.", "Move toward nonce/hash-based CSP and reduce inline script/style allowances after UI compatibility testing."],
            ["M-02", "MFA secrets and high-sensitivity form data are not evidenced as field-level encrypted.", "Medium", "Database compromise impact remains relevant; exploitability is constrained by DB and app boundaries.", "DB boundary, RBAC, hashed passwords/tokens and environment-managed secrets.", "Apply KMS/Key Vault envelope encryption for MFA secrets and classified PII fields."],
            ["M-03", "Upload security does not yet include a full malware scanning/quarantine release service.", "Medium", "Validation blocks common malformed/polyglot cases but cannot fully replace AV/sandbox scanning for uploaded content.", "Upload rate limit, size/MIME/ext/magic-byte validation, private storage default and presigned quarantine path.", "Integrate AV scanning and approval workflow before public or downstream file use."],
            ["M-04", "Outbound form dispatch depends on correct allowlist and egress governance.", "Medium", "SSRF risk is materially reduced but misconfigured allowlists or DNS/egress gaps can still cause exposure.", "HTTPS required, private/localhost hosts blocked and FORM_DISPATCH_ALLOWED_HOSTS required in production.", "Use egress firewall/proxy and DNS pinning or controlled resolver policy."],
            ["M-05", "Residual moderate dependency advisories remain where no safe non-breaking fix is available in current version line.", "Medium", "No critical/high audit findings remain; moderate advisories require lifecycle tracking.", "npm audit fixes applied; high/critical eliminated across reviewed packages.", "Track Next.js/PostCSS advisory remediation and retest when safe upstream versions are available."],
            ["L-01", "Detailed health endpoints could expose operational data if deployment toggles are misconfigured.", "Low", "Production protection is implemented; residual is operational misconfiguration.", "Health token/private internal diagnostic checks.", "Restrict at ingress/network layer."],
            ["L-02", "Admin route permission drift may recur as new routes are added.", "Low", "Current reviewed gaps remediated; future code changes could introduce drift.", "RBAC middleware and Prisma role/permission model.", "Add automated route permission coverage tests."],
            ["L-03", "Filemanager API key is standalone and separate from core identity.", "Low", "Production fail-closed reduces exposure, but key lifecycle must be managed.", "API_KEY required in production and CORS wildcard rejected.", "Integrate with central IAM/API gateway for production."],
            ["L-04", "Public stock API may be abused at low volume.", "Low", "Rate limits and validation constrain abuse.", "Per-IP limiter and symbol/date/interval validation.", "Move high-volume controls to WAF/edge."],
            ["L-05", "Presigned upload flow may be enabled by deployment choice.", "Low", "Default production disabled; if enabled, validation, short TTL and quarantine lower impact.", "PRESIGNED_UPLOAD_ENABLED=false default and quarantine prefix.", "Require security review before enabling in production."],
            ["L-06", "Object storage can be made public by explicit env opt-in.", "Low", "Default is private; residual is deployment governance.", "S3/Azure public settings require explicit true values.", "Cloud posture management and block-public-access policies."],
            ["L-07", "Audit logs may require stronger immutability for compliance evidence.", "Low", "Application logging exists; immutable log sink not evidenced.", "Activity logging and redacted metadata.", "Ship logs to SIEM/WORM or immutable storage."],
        ],
        [0.6, 1.7, 0.8, 2.0, 1.7, 1.9],
        GREEN,
    )

    add_heading(doc, "14. Mitigation and Hardening Recommendations", 1)
    styled_table(
        doc,
        ["Priority", "Recommendation", "Rationale", "Owner"],
        [
            ["P1", "Mandate MFA for privileged/admin roles.", "MFA exists, but global mandatory enforcement was not found. This further reduces account takeover risk.", "Application / Security"],
            ["P1", "Deploy malware scanning and quarantine release workflow for uploads.", "Current validation is strong but not equivalent to AV/sandbox scanning.", "Application / Platform"],
            ["P1", "Adopt managed field-level encryption for MFA secrets and classified PII.", "Reduces database-compromise blast radius.", "Application / Platform"],
            ["P2", "Add CI secret scanning, dependency scanning and route-permission regression tests.", "Prevents recurrence of source-control secret exposure, dependency drift and RBAC gaps.", "DevSecOps"],
            ["P2", "Enforce egress proxy/firewall for outbound dispatch and external integrations.", "Complements application allowlist and mitigates SSRF bypass/misconfiguration paths.", "Platform / Network"],
            ["P2", "Move CSP to nonce/hash-based policy after compatibility testing.", "Reduces browser injection blast radius without breaking Next.js/admin UI behavior.", "Frontend / Security"],
            ["P3", "Forward audit logs to immutable SIEM/WORM storage.", "Improves repudiation evidence and audit readiness.", "Security Operations"],
        ],
        [0.8, 2.4, 2.7, 1.5],
        BLUE,
    )

    add_heading(doc, "15. Verification Evidence", 1)
    styled_table(
        doc,
        ["Verification", "Result", "Evidence Summary"],
        [
            ["backend npm run build", "Passed", "TypeScript backend compiled after remediation."],
            ["frontend npm run build", "Passed", "Next.js admin frontend built after cookie-only session/CSP changes."],
            ["web npm run build", "Passed", "Public web application built after stock API rate-limit validation changes."],
            ["filemanager npm run build", "Passed", "Standalone filemanager built after production fail-closed auth/CORS changes."],
            ["npm audit --omit=dev - backend", "No Critical/High", "One Low advisory remained in diff dependency path after safe audit fixes."],
            ["npm audit --omit=dev - frontend", "No Critical/High", "Moderate Next/PostCSS advisories remain without safe non-breaking fix in current line."],
            ["npm audit --omit=dev - web", "No Critical/High", "Moderate Next/PostCSS-related advisories remain without safe non-breaking fix in current line."],
            ["npm audit --omit=dev - filemanager", "0 vulnerabilities", "Audit clean after npm audit fix."],
            ["Git secret posture", "Improved", "backend/.env and frontend/.env removed from Git index; nested .env ignore rules added."],
        ],
        [2.0, 1.4, 4.0],
        TEAL,
    )

    add_heading(doc, "Appendix A. Evidence File References", 1)
    styled_table(
        doc,
        ["Area", "Representative Files Reviewed / Updated"],
        [
            ["Auth/session", "backend/src/controllers/auth.controller.ts; backend/src/controllers/mfa.controller.ts; backend/src/utils/authResponse.util.ts; frontend/src/context/AuthContext.tsx; frontend/src/services/base.service.ts; frontend/src/services/baseCrud.service.ts; frontend/src/services/auth.service.ts"],
            ["Authorization/RBAC", "backend/src/routes/cms/page.routes.ts; backend/src/routes/dashboard.routes.ts; backend/src/routes/analytics.routes.ts; backend/src/middleware/auth.middleware.ts"],
            ["Upload/storage", "backend/src/middleware/upload.middleware.ts; backend/src/routes/upload.routes.ts; backend/src/controllers/upload.controller.ts; backend/src/routes/filemanager.routes.ts; backend/src/modules/form-modules/formModule.routes.ts; backend/src/services/s3/s3Service.ts; backend/src/services/azureStorage.service.ts"],
            ["External integrations", "backend/src/modules/form-modules/formSubmissionDispatch.service.ts; backend/src/controllers/linknetEnterprise.controller.ts; backend/src/controllers/linknetMedia.controller.ts; web/app/api/stock/quote/route.ts; web/app/api/stock/historical/route.ts; web/lib/apiRateLimit.ts"],
            ["Operational/config", ".gitignore; backend/.env.example; backend/src/middleware/environmentValidator.ts; backend/src/routes/health.routes.ts; frontend/next.config.ts; backend/k8s/*"],
            ["Data model", "backend/prisma/schema.prisma"],
        ],
        [2.0, 5.4],
    )

    add_heading(doc, "Appendix B. Glossary", 1)
    styled_table(
        doc,
        ["Term", "Definition"],
        [
            ["STRIDE", "Threat modeling categories: Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service and Elevation of Privilege."],
            ["DFD", "Data Flow Diagram; visual model of entities, processes, data stores and trust boundaries."],
            ["RBAC", "Role-Based Access Control using roles and permissions to authorize application actions."],
            ["CSRF", "Cross-Site Request Forgery; attack where authenticated browser state is abused by a malicious site."],
            ["HttpOnly cookie", "Cookie inaccessible to JavaScript, reducing token theft through XSS."],
            ["Presigned upload", "Temporary direct object-storage upload URL. In this implementation it is disabled by default in production and quarantined if enabled."],
            ["KMS / Key Vault", "Managed key service used to protect secrets and encryption keys."],
        ],
        [1.8, 5.6],
    )

    target = OUT
    try:
        doc.save(target)
    except PermissionError:
        target = FALLBACK_OUT
        doc.save(target)

    with zipfile.ZipFile(target) as package:
        media = [name for name in package.namelist() if name.startswith("word/media/")]
    print(f"Generated: {target}")
    print(f"Diagrams embedded: {len(media)}")


if __name__ == "__main__":
    build_document()
